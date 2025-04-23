from docx import Document

from convert_markdown.utils.markdown_utils import remove_citation_tokens, add_markdown_runs


def test_remove_citation_tokens_strips_citations_and_normalizes_spacing():
    raw = "Hello citeturn1 world .  Next."
    cleaned = remove_citation_tokens(raw)
    assert cleaned == "Hello world. Next."

    raw2 = "Test foo123.   Multiple   spaces ."
    cleaned2 = remove_citation_tokens(raw2)
    assert cleaned2 == "Test. Multiple spaces."


def test_add_markdown_runs_bold_and_plain_runs():
    doc = Document()
    p = doc.add_paragraph()
    add_markdown_runs(p, "**Bold** and normal text")
    runs = p.runs
    assert len(runs) == 2
    assert runs[0].text == "Bold"
    assert runs[0].bold is True
    assert runs[1].text == " and normal text"
    assert not runs[1].bold


def test_add_markdown_runs_multiple_bold_segments():
    doc = Document()
    p = doc.add_paragraph()
    add_markdown_runs(p, "Start **one** middle **two** end")
    runs = p.runs
    texts = [r.text for r in runs]
    flags = [bool(r.bold) for r in runs]
    assert texts == ["Start ", "one", " middle ", "two", " end"]
    assert flags == [False, True, False, True, False]