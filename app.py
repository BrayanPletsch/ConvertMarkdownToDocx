import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

markdown_text = """
(Insira seu texto em markdown aqui)
"""

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

normal_style = doc.styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(12)
normal_style.font.color.rgb = RGBColor(0, 0, 0)
normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

heading_levels = {1: 18, 2: 16, 3: 14}
for level, size in heading_levels.items():
    style = doc.styles[f'Heading {level}']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.font.bold = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

for section in doc.sections:
    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = "PAGE"
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

if len(doc.paragraphs) > 0:
    doc.add_page_break()

toc_title_paragraph = doc.add_paragraph("SUMÁRIO")
toc_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_title_paragraph.runs[0].bold = True

toc_paragraph = doc.add_paragraph()
run = toc_paragraph.add_run()
fld_begin = OxmlElement('w:fldChar')
fld_begin.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'TOC \\o "1-3" \\h \\z \\u'
fld_separate = OxmlElement('w:fldChar')
fld_separate.set(qn('w:fldCharType'), 'separate')

fld_separator_text = OxmlElement('w:t')
fld_separator_text.text = "Atualize o campo para gerar o sumário.(F9 ou clique direito → Atualizar campo)"
fld_separate.append(fld_separator_text)
fld_end = OxmlElement('w:fldChar')
fld_end.set(qn('w:fldCharType'), 'end')

run._r.append(fld_begin)
run._r.append(instr)
run._r.append(fld_separate)
run._r.append(fld_end)

doc.add_page_break()

lines = markdown_text.splitlines()
i = 0
while i < len(lines):
    line = lines[i]
    if line.strip().startswith('#'):
        level = 0
        while level < len(line) and line[level] == '#':
            level += 1
        heading_text = line[level:].strip()
        if heading_text:
            doc.add_heading(heading_text, level=level)
        i += 1
        continue

    if '|' in line and i + 1 < len(lines) and re.match(r'^\s*[\|\:\-\s]+\s*$', lines[i + 1]):
        header_line = line
        j = i + 2
        body_lines = []
        while j < len(lines) and '|' in lines[j]:
            body_lines.append(lines[j])
            j += 1

        def split_cells(row):
            parts = row.strip().strip('|').split('|')
            return [cell_text.strip() for cell_text in parts]
        header_cells = split_cells(header_line)
        col_count = len(header_cells)
        row_count = len(body_lines)
        table = doc.add_table(rows=row_count + 1, cols=col_count)
        table.style = 'Table Grid'

        for ci, cell_text in enumerate(header_cells):
            cell = table.cell(0, ci)
            cell_para = cell.paragraphs[0]
            parts = re.split(r'(\*\*[^*]+\*\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    cell_para.add_run(part[2:-2]).bold = True
                else:
                    cell_para.add_run(part)

        for ri, body_line in enumerate(body_lines, start=1):
            cells = split_cells(body_line)
            for ci, cell_text in enumerate(cells):
                cell = table.cell(ri, ci)
                cell_para = cell.paragraphs[0]
                parts = re.split(r'(\*\*[^*]+\*\*)', cell_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        cell_para.add_run(part[2:-2]).bold = True
                    else:
                        cell_para.add_run(part)
        i = j
        continue

    if line.strip().startswith('- ') or line.strip().startswith('* '):
        item_text = line.strip()[2:].strip()
        doc.add_paragraph(item_text, style='List Bullet')
        i += 1
        continue

    if re.match(r'^\d+\.\s', line.strip()):
        item_text = re.sub(r'^\d+\.\s*', '', line.strip())
        doc.add_paragraph(item_text, style='List Number')
        i += 1
        continue

    if re.match(r'^\s*[-_*]{3,}\s*$', line):
        hr_para = doc.add_paragraph()
        p = hr_para._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
        i += 1
        continue

    if line.strip() == "":
        doc.add_paragraph("")
        i += 1
        continue

    paragraph = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*)', line)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)
    i += 1

os.makedirs('output', exist_ok=True)
timestamp = datetime.now().strftime("%Y%m%d_%H%M")
base_name = f"relatorio_{timestamp}"
filename = base_name + ".docx"
filepath = os.path.join("output", filename)
counter = 1

while os.path.exists(filepath):
    filename = f"{base_name}_{counter}.docx"
    filepath = os.path.join("output", filename)
    counter += 1

try:
    doc.save(filepath)
    print(f"Documento salvo com sucesso: {filepath}")
except Exception as e:
    print(f"Falha ao salvar o documento: {e}")